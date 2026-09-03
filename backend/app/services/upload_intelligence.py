import io
import os
from typing import Tuple


MAX_UPLOAD_BYTES = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".png", ".jpg", ".jpeg", ".txt", ".eml"}


def validate_upload(file_name: str, content: bytes) -> str:
    extension = os.path.splitext(os.path.basename(file_name))[1].lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise ValueError("Unsupported file type. Use PDF, DOCX, TXT, EML, JPG, or PNG.")
    if not content:
        raise ValueError("The selected file is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise ValueError("The selected file is larger than the 10 MB upload limit.")
    return extension


def extract_upload_text(file_name: str, content: bytes) -> Tuple[str, str]:
    """Return extracted text and the extraction mode used; empty text is a safe fallback."""
    extension = os.path.splitext(os.path.basename(file_name))[1].lower()
    try:
        if extension in {".txt", ".eml"}:
            return content.decode("utf-8", errors="replace")[:50000], "plain_text"
        if extension == ".docx":
            from docx import Document

            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
            return "\n".join(paragraphs)[:50000], "docx_text"
        if extension == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            return "\n".join((page.extract_text() or "") for page in reader.pages)[:50000], "pdf_text"
    except Exception:
        return "", "extraction_unavailable"
    return "", "metadata_only"
