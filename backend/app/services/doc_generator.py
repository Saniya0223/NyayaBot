import os
import uuid
from datetime import datetime
from html.parser import HTMLParser

from docx import Document
from jinja2 import Environment, FileSystemLoader
from xhtml2pdf import pisa

from app.config import settings
from app.schemas.document import DocumentResponse
from app.schemas.fact_graph import FactGraphSchema


class _PlainTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []

    def handle_data(self, data):
        value = data.strip()
        if value:
            self.parts.append(value)

    def get_text(self):
        return "\n".join(self.parts)


class DocumentGenerator:
    """Render deterministic templates from confirmed facts into HTML, PDF, and DOCX."""

    def __init__(self):
        self.env = Environment(loader=FileSystemLoader(settings.TEMPLATES_DIR))

    def generate_document(
        self,
        case_id: str,
        doc_type: str,
        fact_graph: FactGraphSchema,
        appropriate_forum: str = "Competent authority",
        custom_data: dict = None,
    ) -> DocumentResponse:
        custom_data = custom_data or {}
        now = datetime.now()
        date_today = now.strftime("%d-%m-%Y")
        notice_ref = f"NYA/{now.year}/{case_id[:8].upper()}"

        template_map = {
            "GENERAL_COMPLAINT_LETTER": ("general_complaint_letter.html", f"General Complaint Letter - {fact_graph.opposite_party.name}"),
            "FORMAL_LEGAL_NOTICE": ("consumer_legal_notice.html", f"Consumer Grievance Letter - {fact_graph.opposite_party.name}"),
            "EDAAKHIL_COMPLAINT": ("edaakhil_consumer_complaint.html", f"Consumer Complaint Draft - {fact_graph.opposite_party.name}"),
            "TENANT_DEMAND_NOTICE": ("tenant_deposit_notice.html", f"Security Deposit Demand Letter - {fact_graph.opposite_party.name}"),
            "SALARY_DEMAND_NOTICE": ("salary_demand_notice.html", f"Salary Demand Letter - {fact_graph.opposite_party.name}"),
            "POLICE_COMPLAINT_BNSS": ("police_complaint_bnss.html", f"Written Police Complaint Draft - {fact_graph.opposite_party.name}"),
            "CYBERCRIME_BANK_FREEZE": ("cybercrime_bank_freeze.html", f"Cyber Fraud Complaint Draft - {fact_graph.opposite_party.name}"),
            "RTI_SEC6": ("rti_application_sec6.html", f"RTI Application Draft - {fact_graph.opposite_party.name}"),
        }
        if doc_type not in template_map:
            doc_type = "GENERAL_COMPLAINT_LETTER"

        template_name, doc_title = template_map[doc_type]
        complainant = fact_graph.complainant.model_dump()
        opposite_party = fact_graph.opposite_party.model_dump()
        financials = fact_graph.financials.model_dump()

        complainant["name"] = custom_data.get("complainant_name") or complainant.get("name")
        complainant["city"] = custom_data.get("complainant_city") or complainant.get("city")
        complainant["address"] = custom_data.get("complainant_address") or complainant.get("address")
        opposite_party["name"] = (
            custom_data.get("opposite_party_name")
            or custom_data.get("recipient_name")
            or custom_data.get("bank_name")
            or custom_data.get("police_station_name")
            or opposite_party.get("name")
        )
        opposite_party["address"] = (
            custom_data.get("opposite_party_address")
            or custom_data.get("landlord_address")
            or custom_data.get("property_address")
            or opposite_party.get("address")
        )
        opposite_party["city"] = custom_data.get("complainant_city") or opposite_party.get("city")

        disputed_amount = custom_data.get("disputed_amount")
        if disputed_amount is not None:
            amount = float(disputed_amount)
            financials["amount_paid"] = amount
            financials["refund_claimed"] = amount
            financials["total_claim_amount"] = amount

        context = {
            "complainant": complainant,
            "opposite_party": opposite_party,
            "incident_narrative": custom_data.get("incident_narrative") or fact_graph.incident_narrative,
            "incident_date": custom_data.get("incident_date") or custom_data.get("vacating_date") or fact_graph.incident_date,
            "financials": financials,
            "evidence_inventory": [e.model_dump() for e in fact_graph.evidence_inventory],
            "appropriate_forum": appropriate_forum,
            "date_today": date_today,
            "year_current": str(now.year),
            "notice_ref": notice_ref,
            **custom_data,
        }
        rendered_html = self.env.get_template(template_name).render(**context)

        doc_id = str(uuid.uuid4())
        pdf_filename = f"{doc_type.lower()}_{case_id[:8]}_{doc_id[:6]}.pdf"
        pdf_path = os.path.join(settings.STORAGE_DIR, "documents", pdf_filename)
        with open(pdf_path, "wb") as pdf_file:
            pisa_status = pisa.CreatePDF(rendered_html, dest=pdf_file)
        pdf_download_url = f"/api/v1/documents/download/{pdf_filename}" if not pisa_status.err else None

        docx_filename = f"{doc_type.lower()}_{case_id[:8]}_{doc_id[:6]}.docx"
        docx_path = os.path.join(settings.STORAGE_DIR, "documents", docx_filename)
        parser = _PlainTextExtractor()
        parser.feed(rendered_html)
        word_document = Document()
        word_document.add_heading(doc_title, level=1)
        for line in parser.get_text().splitlines():
            word_document.add_paragraph(line)
        word_document.save(docx_path)

        citation_map = {
            "FORMAL_LEGAL_NOTICE": [
                {"act": "Consumer Protection Act, 2019", "section": "Section 2(11)", "title": "Deficiency in service"},
                {"act": "Consumer Protection Act, 2019", "section": "Section 2(47)", "title": "Unfair trade practice"},
            ],
            "EDAAKHIL_COMPLAINT": [
                {"act": "Consumer Protection Act, 2019", "section": "Section 35", "title": "Manner in which complaint shall be made"},
            ],
            "TENANT_DEMAND_NOTICE": [
                {"act": "Indian Contract Act, 1872", "section": "Section 73", "title": "Compensation for breach of contract"},
            ],
            "SALARY_DEMAND_NOTICE": [
                {"act": "Applicable employment and wage law", "section": "State/fact specific", "title": "Payment of earned wages"},
            ],
            "POLICE_COMPLAINT_BNSS": [
                {"act": "Bharatiya Nagarik Suraksha Sanhita, 2023", "section": "Section 173", "title": "Information in cognizable cases"},
            ],
            "CYBERCRIME_BANK_FREEZE": [
                {"act": "Information Technology Act, 2000", "section": "Section 66D", "title": "Cheating by personation using computer resource"},
                {"act": "RBI/2017-18/15", "section": "Paragraphs 6-10", "title": "Customer liability for unauthorised electronic transactions"},
            ],
            "RTI_SEC6": [
                {"act": "Right to Information Act, 2005", "section": "Section 6(1)", "title": "Request for obtaining information"},
            ],
        }

        annexures = [
            {"label": ev.annexure_label or f"Annexure A-{index + 1}", "name": ev.doc_name}
            for index, ev in enumerate(fact_graph.evidence_inventory)
        ]
        return DocumentResponse(
            id=doc_id,
            case_id=case_id,
            doc_type=doc_type,
            title=doc_title,
            content_html=rendered_html,
            pdf_download_url=pdf_download_url,
            docx_download_url=f"/api/v1/documents/download/{docx_filename}",
            statutory_citations=citation_map.get(doc_type, []),
            annexures=annexures,
            created_at=datetime.utcnow().isoformat(),
        )


doc_generator = DocumentGenerator()
